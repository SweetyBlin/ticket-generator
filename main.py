from docx import Document
from docx.shared import Mm
import random

#
theory_count = 11

document = Document()
section = document.sections[0]
section.left_margin = Mm(15)
section.right_margin = Mm(15)
section.top_margin = Mm(15)
section.bottom_margin = Mm(10)

theory = open('questions.txt', 'r+', encoding="utf-8")
practice = open('practice.txt', 'r+', encoding="utf-8")
theory_clear = [line.rstrip() for line in theory.readlines()]
practice_clear = [line.rstrip() for line in practice.readlines()]
random.shuffle(theory_clear)

for i in range(0, len(theory_clear), theory_count):
    p = document.add_paragraph(f"Билет №{int(i/theory_count+1)}")
    p.alignment = 1
    document.add_paragraph(f"Теоретические вопросы:")
    for count in range(theory_count):
        p = document.add_paragraph(f"{count + 1}. {theory_clear[i + count]}")
        p.alignment = 3
    document.add_paragraph(f"Практическое задание:")
    p = document.add_paragraph(f"{random.choice(practice_clear)}")
    p.alignment = 3
    document.add_paragraph('\n')


document.save('test.docx')
